import numpy as np
import cv2
from numpy import array
from docx import Document

# Membaca gambar menggunakan OpenCV
# Menggantikan "gambar.jpg" dengan nama file gambar Anda
image = cv2.imread(
    r"/home/faridwijayaaa/Documents/img/unnamed.jpg", cv2.IMREAD_COLOR)

# Mengubah gambar menjadi matriks NumPy
matrix = np.array(image)

# Menyusun matriks menjadi ukuran yang diinginkan
desired_rows = 175
desired_cols = 200
matrix = cv2.resize(matrix, (desired_cols, desired_rows))

# Menyimpan matriks ke dalam file gambar
output_jpg = "output.jpg"  # Nama file output yang diinginkan
cv2.imwrite(output_jpg, matrix)

output_file = "output.docx"  # Nama file output yang diinginkan
doc = Document()
doc.add_paragraph("Hasil Matriks:")
for row in matrix:
    doc.add_paragraph(" ".join(str(cell) for cell in row))
doc.save(output_file)


# Mengecek bentuk matriks hasil
print("Bentuk matriks: ", matrix.shape)
print("Matriks berhasil disimpan sebagai file .docx:", output_file)
print("Matriks berhasil disimpan sebagai file gambar:", output_file)
